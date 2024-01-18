import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_LINE_SPACING

import os

from docx2pdf import convert

from addHyperlink import add_hyperlink

files_directory=r"C:\Users\simon\Desktop\mycoolshit\SeleniumCoop\venvForSeleniumCoop\TailoredCoverLetters3"
# paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
allTxtFiles=os.listdir(files_directory)
new_txt_file_list = [filename.replace(".txt", "") for filename in allTxtFiles] 

for filename in new_txt_file_list: #get list of files from directory
    print(filename)

header=""#declare header variable
desc=""#declare decs variable for body of cover letter

#document=Document() 



for eachFile in os.listdir(files_directory):
    # Construct the full path to each file using os.path.join
    jobTitle=eachFile.replace(".txt","")
    print("Job Title:")
    eFile=os.path.join(files_directory,eachFile)
    with open(eFile,"r") as file:
        for line_number,line in enumerate(file, 1):  #enumerare just indexs the text document line by line so we can organize it from that
            #organize line by line # user range loop, assumes all files are same length tho
            print(f"Line {line_number}: {line.strip()}")  # Stripping newline characters for better output
            if line_number == 1:
                name=line
            elif line_number == 2 or line_number == 3 or line_number == 4:
                header+=line
            else:
                desc+=line

    #print (desc)

    #strip witespace to fix formating error latter
    name=name.strip("\n")
    header=header.strip("\n")
    #create new document
    document=Document()

    #set font and some margins 
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    #set margin -------
    sections=document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)


    # make name variable
    # header variable
    # full desc variable
    #add variable to paragraphs and hope thats formated
    #----------Adds Name and styles it
    paragraph1 = document.add_paragraph()
    paragraph1.add_run(name)
    paragraph1_format = paragraph1.paragraph_format
    #paragraph1_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    paragraph1_format.space_before = Pt(0)
    paragraph1_format.space_after = Pt(0)
    paragraph1_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph1.runs[0].font.size = Pt(14)
    paragraph1.runs[0].bold= True


 
    #------ Adds contact information and styles it, adds hyperlinks, 
    paragraph2 = document.add_paragraph()

    #------- Add hyperlinks
    # start by partitioning links into runs
    # listOfContactInfo=header.partition("https://www.linkedin.com/in/simon-amable-59ab091ab/") #index 1 is our hyperlink
    # for cInfo in listOfContactInfo:
    #     cInfo=cInfo.strip(", ",)
    # for cInfo in listOfContactInfo:
    #     print(cInfo)
    # print("CONTACT INFO TUPLE")
    # print(listOfContactInfo)
    #partition and dont add last
    contactPartitioned=header.partition("simonamable@gmail.com")
    paragraph2.add_run(contactPartitioned[0])
    paragraph2.add_run(contactPartitioned[1]+"\n")


    add_hyperlink(paragraph2, 'https://www.linkedin.com/in/simon-amable-59ab091ab/', 'https://www.linkedin.com/in/simon-amable-59ab091ab/')
    paragraph2.add_run(" | ")

    add_hyperlink(paragraph2, 'https://github.com/SimonAmable', 'https://github.com/SimonAmable')

    
    
    #contact info formatting
    paragraph2_format = paragraph2.paragraph_format
    paragraph2_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph2_format.space_before = Pt(0)  # Remove the space before the paragraph
    paragraph2_format.space_after = Pt(0)
    #could modify paragraph font instead
    paragraph2.style.font.size = Pt(13)
    # paragraph2.runs[0].font.size = Pt(12)
    # paragraph2.runs[1].font.size = Pt(12)

    

    #----------Adds Body and styles it
    paragraph3 = document.add_paragraph()
    #seperate body into 3 runs so i can bold 1, use rpartition
    #BodyInList[1] will be Bolded 
    BodyInList=desc.partition("This position was looking for someone enthusiastic about programming, eager to learn new technologies, or who has strong transferable problem-solving skills so I wanted to give you a clear demonstration of why thats me by developing a completely automated program to get this application to you")
    paragraph3.add_run(BodyInList[0])
    paragraph3.add_run(BodyInList[1]) #BodyInList[1] will be Bolded, you can select one line from body
    paragraph3.add_run(BodyInList[2])



    paragraph3_format = paragraph3.paragraph_format
    paragraph3_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph3_format.space_before = Pt(0)
    paragraph3_format.space_after = Pt(0)
    paragraph3.runs[0].font.size = Pt(12)
    paragraph3.runs[1].font.size = Pt(12)
    paragraph3.runs[2].font.size = Pt(12)

    paragraph3.runs[1].bold= True
    



    #construct output file path and log in terminal
    output_directory2 = r"C:\Users\simon\Desktop\mycoolshit\SeleniumCoop\venvForSeleniumCoop\FinishedCoverLetters4"
    #jobTitle=jobTitle.replace(' ', '_').replace('/', '').strip('\n')
    print ("job title: "+jobTitle)
    output_filename2 = jobTitle.replace("modified_", "") + ".docx"
    
    output_filepath2 = os.path.join(output_directory2, output_filename2)
    print ("saving file to: " +output_filepath2)
    document.save(output_filepath2)
    
    #convert to files to pdf with seprate script!!!!!!!!!!!!!!!!!!





    #pdf_output_path = os.path.join(output_directory2, (output_filename2.replace(".docx", ".pdf").replace("modified_","")  ) )

    

     # Reset the variables for the next file
    name = ""
    jobTitle = ""
    header = ""
    desc = ""

    #make document create a function, take in the name for the and file path?


from docx2pdf import convert#convert files to pdf
convert(output_directory2)